"""
CV -> structured fields.

Two stages, deliberately separated:

  1. Get text off the page. Boring, deterministic, free. PyMuPDF for PDF,
     python-docx for Word. If a PDF yields almost no text it is a scan, and we
     fall back to sending page images to the model instead.

  2. Turn that text into fields. This is the only stage a model touches, and it
     is constrained by a response schema so the output is a typed object rather
     than prose we have to parse and pray over. The previous attempt at a
     product here died on output quality; the defence is that every field is
     checkable against the source document by eye in about four seconds.

The model is told, repeatedly, not to invent. A missing phone number must come
back null. An invented one is worse than a blank, because a recruiter will send
it to a client.
"""

import base64
import json
import os
import re
import time

import requests

MODEL = "gemini-3.5-flash-lite"
ENDPOINT = f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL}:generateContent"

# Enough text that we believe the PDF has a real text layer rather than being a
# scan with a stray caption on it.
TEXT_LAYER_MIN_CHARS = 250


# --------------------------------------------------------------------- schema
# Gemini enforces this, so the caller always gets these keys with these types.
# Anything the CV does not state comes back null or empty, never guessed.
SCHEMA = {
    "type": "object",
    "properties": {
        "name": {"type": "string", "nullable": True},
        "headline": {"type": "string", "nullable": True},
        "location": {"type": "string", "nullable": True},
        "email": {"type": "string", "nullable": True},
        "phone": {"type": "string", "nullable": True},
        "links": {"type": "array", "items": {"type": "string"}},
        "summary": {"type": "string", "nullable": True},
        "experience": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "employer": {"type": "string", "nullable": True},
                    "title": {"type": "string", "nullable": True},
                    "location": {"type": "string", "nullable": True},
                    "start": {"type": "string", "nullable": True},
                    "end": {"type": "string", "nullable": True},
                    "current": {"type": "boolean"},
                    "bullets": {"type": "array", "items": {"type": "string"}},
                },
                "required": ["employer", "title", "start", "end", "current", "bullets"],
            },
        },
        "education": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "institution": {"type": "string", "nullable": True},
                    "qualification": {"type": "string", "nullable": True},
                    "start": {"type": "string", "nullable": True},
                    "end": {"type": "string", "nullable": True},
                    "detail": {"type": "string", "nullable": True},
                },
                "required": ["institution", "qualification", "start", "end"],
            },
        },
        "skills": {"type": "array", "items": {"type": "string"}},
        "languages": {"type": "array", "items": {"type": "string"}},
        "certifications": {"type": "array", "items": {"type": "string"}},
    },
    "required": ["name", "experience", "education", "skills"],
}


INSTRUCTIONS = """You are extracting structured data from a candidate CV for a recruitment agency.

Rules, in order of importance:

1. NEVER invent information. If the CV does not state something, return null for
   that field, or an empty array. A blank field is correct; a plausible guess is
   a serious error, because a recruiter will forward it to a client as fact.

2. Copy wording from the CV. Do not rewrite, improve, summarise or expand
   bullets. Fix only obvious extraction damage: a word split across a line
   break, a hyphen inserted by wrapping, doubled spaces.

3. Text extracted from a PDF often breaks things across lines. Reassemble them.
   An email may arrive as two fragments on separate lines; join them with no
   space. The same applies to URLs.

4. Ignore page furniture: headers, footers, page numbers, "confidential"
   notices, and any text the candidate did not write as CV content.

5. Dates: normalise to YYYY-MM where the month is known, otherwise YYYY. If a
   role is ongoing ("Present", "Current", "to date"), set end to null and
   current to true. Otherwise current is false.

6. Sidebars are content. Skills, languages and contact details often live in a
   separate column and must still be captured.

7. Preserve the order roles appear in, which is normally most recent first.

Return only the structured object."""


# ------------------------------------------------------------------- reading
def read_pdf_text(path):
    import pymupdf

    doc = pymupdf.open(path)
    pages = [p.get_text("text") for p in doc]
    doc.close()
    return "\n".join(pages).strip()


def read_pdf_images(path, max_pages=4, dpi=140):
    """Fallback for scanned CVs: hand the model the page as an image."""
    import pymupdf

    doc = pymupdf.open(path)
    out = []
    for page in list(doc)[:max_pages]:
        pix = page.get_pixmap(dpi=dpi)
        out.append(base64.b64encode(pix.tobytes("png")).decode())
    doc.close()
    return out


def read_docx_text(path):
    import docx

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    # Tables are where Word CVs hide their layout, and their text is invisible
    # to a paragraph-only walk.
    for table in d.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(x for x in parts if x and x.strip()).strip()


def read_document(path):
    """Returns (text, images). Exactly one of them is populated."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx_text(path), []
    if ext == ".pdf":
        text = read_pdf_text(path)
        if len(re.sub(r"\s", "", text)) >= TEXT_LAYER_MIN_CHARS:
            return text, []
        return "", read_pdf_images(path)
    if ext in (".txt", ".md"):
        with open(path, encoding="utf-8", errors="replace") as fh:
            return fh.read().strip(), []
    raise ValueError(f"unsupported file type: {ext}")


# ----------------------------------------------------------------- the model
def _call(parts, api_key, timeout=90):
    body = {
        "contents": [{"parts": parts}],
        "systemInstruction": {"parts": [{"text": INSTRUCTIONS}]},
        "generationConfig": {
            "temperature": 0,
            "responseMimeType": "application/json",
            "responseSchema": SCHEMA,
            "maxOutputTokens": 8192,
        },
    }
    res = requests.post(
        ENDPOINT, params={"key": api_key},
        json=body, timeout=timeout,
        headers={"content-type": "application/json"},
    )
    if res.status_code != 200:
        raise RuntimeError(f"gemini {res.status_code}: {res.text[:400]}")
    return res.json()


def extract(path, api_key=None, retries=2):
    api_key = api_key or os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY is not set")

    text, images = read_document(path)
    if not text and not images:
        raise RuntimeError(f"no readable content in {path}")

    if text:
        parts = [{"text": f"CV text:\n\n{text}"}]
        source = "text"
    else:
        parts = [{"text": "CV pages follow as images."}] + [
            {"inlineData": {"mimeType": "image/png", "data": b64}} for b64 in images
        ]
        source = "image"

    last = None
    for attempt in range(retries + 1):
        try:
            payload = _call(parts, api_key)
            raw = payload["candidates"][0]["content"]["parts"][0]["text"]
            data = json.loads(raw)
            data["_source"] = source
            data["_usage"] = payload.get("usageMetadata", {})
            return data
        except Exception as e:  # noqa: BLE001 - retry anything transient
            last = e
            if attempt < retries:
                time.sleep(1.5 * (attempt + 1))
    raise RuntimeError(f"extraction failed after {retries + 1} attempts: {last}")


if __name__ == "__main__":
    import sys

    target = sys.argv[1] if len(sys.argv) > 1 else "samples/messy-cv.pdf"
    result = extract(target)
    print(json.dumps(result, indent=2, ensure_ascii=False))
