"""
Structured fields -> the agency's branded CV.

The single most important behaviour in this file is what it leaves out.

An agency reformats a CV for one commercial reason above all others: to remove
the candidate's name and contact details before the CV goes to a client. If the
client can read the candidate's email, the client can hire them directly and the
agency loses a fee worth thousands. Everything else here is presentation; this
part is the business.

So redaction is the default, and turning it off has to be deliberate.
"""

import os
import re
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

MONTHS = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun",
          "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]


def pretty_date(value):
    """'2022-03' -> 'Mar 2022'. Anything unexpected passes through untouched."""
    if not value:
        return None
    m = re.fullmatch(r"(\d{4})-(\d{1,2})", value.strip())
    if m:
        year, mon = m.group(1), int(m.group(2))
        if 1 <= mon <= 12:
            return f"{MONTHS[mon]} {year}"
    return value.strip()


def date_range(role):
    start = pretty_date(role.get("start"))
    end = "Present" if role.get("current") else pretty_date(role.get("end"))
    if start and end:
        return f"{start} – {end}"
    return start or end or ""


def initials(name):
    if not name:
        return "Candidate"
    parts = [p for p in re.split(r"\s+", name.strip()) if p]
    return "".join(p[0].upper() for p in parts[:3])


class Brand:
    """One agency's presentation settings."""

    def __init__(self, name, colour="1F4E5F", logo=None, footer=None,
                 contact=None, redact=True):
        self.name = name
        self.colour = RGBColor.from_string(colour.lstrip("#").upper())
        self.logo = logo if logo and os.path.exists(logo) else None
        self.footer = footer
        self.contact = contact
        self.redact = redact


def _shade(paragraph, hex_colour):
    """Paragraph background. python-docx has no API for it, so drop to XML."""
    pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    pr.append(shd)


def _para(doc, text="", size=10, bold=False, italic=False, colour=None,
          space_before=0, space_after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if colour is not None:
            run.font.color.rgb = colour
    return p


def _section_heading(doc, label, brand):
    p = _para(doc, label.upper(), size=9, bold=True, colour=brand.colour,
              space_before=12, space_after=4)
    p.runs[0].font.name = "Calibri"
    # A rule under the heading, again via XML since there's no wrapper for it.
    pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), str(brand.colour))
    borders.append(bottom)
    pr.append(borders)
    return p


def render(data, brand, out_path, reference=None):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ---- masthead ----------------------------------------------------------
    if brand.logo:
        doc.add_picture(brand.logo, width=Inches(1.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        _para(doc, brand.name, size=15, bold=True, colour=brand.colour, space_after=2)

    ref = reference or f"{initials(data.get('name'))}-{date.today():%Y%m}"

    # ---- candidate identity ------------------------------------------------
    # Redacted by default. The reference is how the agency and client talk about
    # this person without the client being able to reach them.
    if brand.redact:
        title = f"Candidate reference {ref}"
    else:
        title = data.get("name") or f"Candidate reference {ref}"

    _para(doc, title, size=18, bold=True, space_before=8, space_after=2)

    if data.get("headline"):
        _para(doc, data["headline"], size=11.5, colour=brand.colour, space_after=2)

    # Location is commercially safe and genuinely useful to a client;
    # email, phone and personal links are not, and never survive redaction.
    facts = []
    if data.get("location"):
        facts.append(data["location"])
    if not brand.redact:
        facts += [x for x in (data.get("email"), data.get("phone")) if x]
        facts += list(data.get("links") or [])
    if facts:
        _para(doc, "  ·  ".join(facts), size=9, colour=RGBColor(0x66, 0x6C, 0x74),
              space_after=2)

    if data.get("summary"):
        _section_heading(doc, "Profile", brand)
        _para(doc, data["summary"], size=10, space_after=2)

    # ---- experience --------------------------------------------------------
    experience = data.get("experience") or []
    if experience:
        _section_heading(doc, "Experience", brand)
        for role in experience:
            line = " — ".join(x for x in (role.get("title"), role.get("employer")) if x)
            _para(doc, line or "Role", size=10.5, bold=True, space_before=7, space_after=0)

            meta = "  ·  ".join(x for x in (date_range(role), role.get("location")) if x)
            if meta:
                _para(doc, meta, size=8.5, italic=True,
                      colour=RGBColor(0x66, 0x6C, 0x74), space_after=3)

            for bullet in role.get("bullets") or []:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Inches(0.22)
                run = p.add_run(bullet)
                run.font.size = Pt(9.5)

    # ---- education ---------------------------------------------------------
    education = data.get("education") or []
    if education:
        _section_heading(doc, "Education", brand)
        for ed in education:
            line = " — ".join(x for x in (ed.get("qualification"), ed.get("institution")) if x)
            _para(doc, line or "Qualification", size=10, bold=True,
                  space_before=5, space_after=0)
            bits = [x for x in (date_range(ed), ed.get("detail")) if x]
            if bits:
                _para(doc, "  ·  ".join(bits), size=8.5, italic=True,
                      colour=RGBColor(0x66, 0x6C, 0x74), space_after=2)

    # ---- lists -------------------------------------------------------------
    for label, key in (("Skills", "skills"), ("Certifications", "certifications"),
                       ("Languages", "languages")):
        values = data.get(key) or []
        if values:
            _section_heading(doc, label, brand)
            _para(doc, "  ·  ".join(values), size=9.5, space_after=2)

    # ---- footer ------------------------------------------------------------
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bits = [brand.footer or brand.name, f"Ref {ref}"]
    if brand.contact:
        bits.insert(1, brand.contact)
    run = footer_p.add_run("  ·  ".join(bits))
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x8A, 0x90, 0x98)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    import json
    import sys

    payload = json.load(open(sys.argv[1], encoding="utf-8"))
    brand = Brand(name="Meridian Talent Partners", colour="1F4E5F",
                  footer="Meridian Talent Partners Ltd",
                  contact="hello@meridiantalent.example")
    out = render(payload, brand, sys.argv[2] if len(sys.argv) > 2 else "out/candidate.docx")
    print(f"wrote {out}")
